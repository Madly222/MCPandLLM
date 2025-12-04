# tools/file_generator_tool.py
import os
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
from io import BytesIO

from openpyxl import Workbook, load_workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, Alignment, Border, Side

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from tools.file_reader_tool import (
    ExtractedContent,
    ExtractedTable,
    ExtractedImage,
    find_file,
    extract_content,
    read_multiple_files,
    EXAMPLES_DIR
)

logger = logging.getLogger(__name__)

BASE_DIR = Path(__file__).resolve().parent.parent
DOWNLOADS_DIR = Path(os.getenv("DOWNLOADS_DIR", BASE_DIR / "downloads"))
SERVER_URL = os.getenv("SERVER_URL", "http://localhost:8000")

DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)


def _generate_filename(base_name: str, extension: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    clean_name = Path(base_name).stem
    return f"{clean_name}_{timestamp}{extension}"


def create_excel(
        contents: List[ExtractedContent],
        output_name: str = "generated",
        title: Optional[str] = None,
        include_images: bool = True,
        separate_sheets: bool = False
) -> Dict[str, Any]:
    try:
        wb = Workbook()
        ws = wb.active

        if title:
            ws.title = title[:31]
        else:
            ws.title = "Данные"

        header_font = Font(bold=True, size=12)
        header_alignment = Alignment(horizontal='center', vertical='center')
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        current_row = 1
        all_images = []

        for content_idx, content in enumerate(contents):
            if separate_sheets and content_idx > 0:
                sheet_name = content.filename[:31] if content.filename else f"Лист_{content_idx + 1}"
                ws = wb.create_sheet(title=sheet_name)
                current_row = 1

            if len(contents) > 1 or content.filename:
                ws.cell(row=current_row, column=1, value=f"Источник: {content.filename}")
                ws.cell(row=current_row, column=1).font = Font(bold=True, italic=True, size=11)
                current_row += 1

            for table in content.tables:
                if table.sheet_name and separate_sheets:
                    ws.cell(row=current_row, column=1, value=f"Таблица: {table.sheet_name}")
                    ws.cell(row=current_row, column=1).font = Font(italic=True)
                    current_row += 1

                for col_idx, header in enumerate(table.headers, 1):
                    cell = ws.cell(row=current_row, column=col_idx, value=header)
                    cell.font = header_font
                    cell.alignment = header_alignment
                    cell.border = thin_border
                current_row += 1

                for row_data in table.rows:
                    for col_idx, value in enumerate(row_data, 1):
                        cell = ws.cell(row=current_row, column=col_idx, value=value)
                        cell.border = thin_border
                    current_row += 1

                current_row += 1

            if content.text and not content.tables:
                lines = content.text.split('\n')
                for line in lines[:100]:
                    if line.strip():
                        ws.cell(row=current_row, column=1, value=line.strip())
                        current_row += 1
                current_row += 1

            if include_images:
                for img in content.images:
                    all_images.append((img, current_row))

        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column].width = adjusted_width

        if include_images and all_images:
            img_sheet = wb.create_sheet(title="Изображения")
            img_row = 1

            for img_data, _ in all_images:
                try:
                    img_stream = BytesIO(img_data.data)
                    xl_img = XLImage(img_stream)

                    xl_img.width = min(xl_img.width, 400)
                    xl_img.height = min(xl_img.height, 300)

                    img_sheet.add_image(xl_img, f"A{img_row}")
                    img_sheet.cell(row=img_row, column=5, value=img_data.filename)

                    img_row += 20
                except Exception as e:
                    logger.warning(f"Ошибка вставки изображения: {e}")

        output_filename = _generate_filename(output_name, ".xlsx")
        output_path = DOWNLOADS_DIR / output_filename

        wb.save(output_path)
        wb.close()

        download_url = f"{SERVER_URL}/download/{output_filename}"

        logger.info(f"Excel создан: {output_path}")

        return {
            "success": True,
            "filename": output_filename,
            "download_url": download_url,
            "path": str(output_path),
            "sources": [c.filename for c in contents],
            "tables_count": sum(len(c.tables) for c in contents),
            "images_count": sum(len(c.images) for c in contents)
        }

    except Exception as e:
        logger.error(f"Ошибка создания Excel: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


def create_word(
        contents: List[ExtractedContent],
        output_name: str = "generated",
        title: Optional[str] = None,
        include_images: bool = True
) -> Dict[str, Any]:
    try:
        doc = Document()

        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        for content_idx, content in enumerate(contents):
            if content.filename:
                doc.add_heading(f"Источник: {content.filename}", level=1)

            if content.text:
                paragraphs = content.text.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        if para_text.startswith('---') and para_text.endswith('---'):
                            doc.add_heading(para_text.strip('- '), level=2)
                        else:
                            doc.add_paragraph(para_text.strip())

            for table_idx, table_data in enumerate(content.tables):
                if table_data.sheet_name:
                    doc.add_heading(f"Таблица: {table_data.sheet_name}", level=2)
                elif len(content.tables) > 1:
                    doc.add_heading(f"Таблица {table_idx + 1}", level=2)

                if table_data.headers or table_data.rows:
                    num_cols = len(table_data.headers) if table_data.headers else len(
                        table_data.rows[0]) if table_data.rows else 0
                    num_rows = (1 if table_data.headers else 0) + len(table_data.rows)

                    if num_cols > 0 and num_rows > 0:
                        doc_table = doc.add_table(rows=num_rows, cols=num_cols)
                        doc_table.style = 'Table Grid'

                        if table_data.headers:
                            header_row = doc_table.rows[0]
                            for idx, header in enumerate(table_data.headers):
                                cell = header_row.cells[idx]
                                cell.text = str(header)
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True

                        start_row = 1 if table_data.headers else 0
                        for row_idx, row_data in enumerate(table_data.rows):
                            doc_row = doc_table.rows[start_row + row_idx]
                            for col_idx, value in enumerate(row_data):
                                if col_idx < num_cols:
                                    doc_row.cells[col_idx].text = str(value)

                        doc.add_paragraph()

            if include_images and content.images:
                doc.add_heading("Изображения", level=2)

                for img in content.images:
                    try:
                        img_stream = BytesIO(img.data)
                        doc.add_picture(img_stream, width=Inches(5))

                        caption = doc.add_paragraph(img.filename)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.runs[0].italic = True

                        doc.add_paragraph()
                    except Exception as e:
                        logger.warning(f"Ошибка вставки изображения: {e}")

            if content_idx < len(contents) - 1:
                doc.add_page_break()

        output_filename = _generate_filename(output_name, ".docx")
        output_path = DOWNLOADS_DIR / output_filename

        doc.save(output_path)

        download_url = f"{SERVER_URL}/download/{output_filename}"

        logger.info(f"Word создан: {output_path}")

        return {
            "success": True,
            "filename": output_filename,
            "download_url": download_url,
            "path": str(output_path),
            "sources": [c.filename for c in contents],
            "tables_count": sum(len(c.tables) for c in contents),
            "images_count": sum(len(c.images) for c in contents)
        }

    except Exception as e:
        logger.error(f"Ошибка создания Word: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


def create_from_template(
        template_name: str,
        contents: List[ExtractedContent],
        output_name: str = "from_template",
        role: Optional[str] = None
) -> Dict[str, Any]:
    template_path = find_file(template_name, role)
    if not template_path:
        template_path = EXAMPLES_DIR / template_name
        if not template_path.exists():
            return {
                "success": False,
                "error": f"Шаблон не найден: {template_name}"
            }

    suffix = template_path.suffix.lower()

    if suffix in ['.xlsx', '.xls']:
        return _create_excel_from_template(template_path, contents, output_name)
    elif suffix == '.docx':
        return _create_word_from_template(template_path, contents, output_name)
    else:
        return {
            "success": False,
            "error": f"Неподдерживаемый формат шаблона: {suffix}"
        }


def _create_excel_from_template(
        template_path: Path,
        contents: List[ExtractedContent],
        output_name: str
) -> Dict[str, Any]:
    try:
        wb = load_workbook(template_path)
        ws = wb.active

        template_structure = _analyze_excel_template(ws)

        current_row = template_structure.get('data_start_row', 2)

        for content in contents:
            for table in content.tables:
                for row_data in table.rows:
                    for col_idx, value in enumerate(row_data, 1):
                        ws.cell(row=current_row, column=col_idx, value=value)
                    current_row += 1

        output_filename = _generate_filename(output_name, ".xlsx")
        output_path = DOWNLOADS_DIR / output_filename

        wb.save(output_path)
        wb.close()

        download_url = f"{SERVER_URL}/download/{output_filename}"

        return {
            "success": True,
            "filename": output_filename,
            "download_url": download_url,
            "template_used": template_path.name,
            "sources": [c.filename for c in contents],
            "tables_count": sum(len(c.tables) for c in contents),
            "images_count": sum(len(c.images) for c in contents)
        }

    except Exception as e:
        logger.error(f"Ошибка создания Excel из шаблона: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


def _create_word_from_template(
        template_path: Path,
        contents: List[ExtractedContent],
        output_name: str
) -> Dict[str, Any]:
    try:
        doc = Document(template_path)

        all_text = "\n\n".join([c.text for c in contents if c.text])
        all_tables = []
        all_images = []

        for content in contents:
            all_tables.extend(content.tables)
            all_images.extend(content.images)

        for paragraph in doc.paragraphs:
            if '{{TEXT}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{TEXT}}', all_text[:5000])
            if '{{SOURCES}}' in paragraph.text:
                sources = ", ".join([c.filename for c in contents])
                paragraph.text = paragraph.text.replace('{{SOURCES}}', sources)
            if '{{DATE}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{DATE}}', datetime.now().strftime("%d.%m.%Y"))

        if all_tables:
            doc.add_heading("Таблицы", level=1)
            for table_data in all_tables:
                if table_data.headers or table_data.rows:
                    num_cols = len(table_data.headers) if table_data.headers else len(table_data.rows[0])
                    num_rows = (1 if table_data.headers else 0) + len(table_data.rows)

                    doc_table = doc.add_table(rows=num_rows, cols=num_cols)
                    doc_table.style = 'Table Grid'

                    if table_data.headers:
                        for idx, header in enumerate(table_data.headers):
                            doc_table.rows[0].cells[idx].text = str(header)

                    start_row = 1 if table_data.headers else 0
                    for row_idx, row_data in enumerate(table_data.rows):
                        for col_idx, value in enumerate(row_data):
                            if col_idx < num_cols:
                                doc_table.rows[start_row + row_idx].cells[col_idx].text = str(value)

                    doc.add_paragraph()

        if all_images:
            doc.add_heading("Изображения", level=1)
            for img in all_images:
                try:
                    img_stream = BytesIO(img.data)
                    doc.add_picture(img_stream, width=Inches(5))
                    doc.add_paragraph()
                except Exception as e:
                    logger.warning(f"Ошибка вставки изображения: {e}")

        output_filename = _generate_filename(output_name, ".docx")
        output_path = DOWNLOADS_DIR / output_filename

        doc.save(output_path)

        download_url = f"{SERVER_URL}/download/{output_filename}"

        return {
            "success": True,
            "filename": output_filename,
            "download_url": download_url,
            "template_used": template_path.name,
            "sources": [c.filename for c in contents],
            "tables_count": sum(len(c.tables) for c in contents),
            "images_count": sum(len(c.images) for c in contents)
        }

    except Exception as e:
        logger.error(f"Ошибка создания Word из шаблона: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


def _analyze_excel_template(ws) -> Dict[str, Any]:
    structure = {
        'has_header': False,
        'header_row': 1,
        'data_start_row': 2,
        'columns': []
    }

    first_row = [cell.value for cell in ws[1]]
    if any(first_row):
        structure['has_header'] = True
        structure['columns'] = [str(v) if v else "" for v in first_row]

    for row_idx in range(1, min(10, ws.max_row + 1)):
        row_values = [cell.value for cell in ws[row_idx]]
        if not any(row_values):
            structure['data_start_row'] = row_idx
            break

    return structure


def generate_file(
        source_files: List[str],
        output_format: str,
        output_name: str = "generated",
        title: Optional[str] = None,
        template_name: Optional[str] = None,
        include_images: bool = True,
        role: Optional[str] = None
) -> Dict[str, Any]:
    contents = read_multiple_files(source_files, role)

    if not contents:
        return {
            "success": False,
            "error": "Не удалось прочитать ни один файл"
        }

    if template_name:
        return create_from_template(template_name, contents, output_name, role)

    output_format = output_format.lower().strip('.')

    if output_format in ['xlsx', 'excel', 'xls']:
        return create_excel(contents, output_name, title, include_images)
    elif output_format in ['docx', 'word', 'doc']:
        return create_word(contents, output_name, title, include_images)
    else:
        return {
            "success": False,
            "error": f"Неподдерживаемый формат: {output_format}. Доступные: xlsx, docx"
        }