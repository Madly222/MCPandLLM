import os
import re
import json
import logging
from pathlib import Path
from typing import Optional, Tuple, List

from agent.memory import memory
from tools.file_tool import try_handle_file_command, select_file
from tools.excel_tool import read_excel, read_excel_for_edit
from tools.search_tool import perform_search, smart_search
from tools.edit_excel_tool import edit_excel, get_excel_preview
from tools.excel_nlu import parse_excel_command
from tools.multi_file_tool import process_multiple_files
from tools.file_reader_tool import get_example_files, find_file
from tools.file_generator_tool import generate_file
from vector_store import vector_store

logger = logging.getLogger(__name__)

BASE_DIR = Path(__file__).resolve().parent.parent
STORAGE_DIR = Path(os.getenv("FILES_DIR", BASE_DIR / "storage"))


def get_role_files_dir(role: str) -> Path:
    return STORAGE_DIR / role


EDIT_TRIGGERS = [
    r"добавь строку",
    r"добавь колонку",
    r"удали строку",
    r"удали колонку",
    r"измени ячейку",
    r"поменяй ячейку",
    r"вставь строку",
    r"новая строка",
    r"новая колонка",
    r"отредактируй",
    r"редактируй",
    r"измени в файле",
    r"измени файл",
    r"обнови файл",
    r"удали.*работ",
    r"удали.*строк",
    r"добавь.*в файл",
    r"добавь.*в таблиц",
]

GENERATE_TRIGGERS = [
    r"создай.*из.*файл",
    r"сделай.*из.*файл",
    r"объедини.*файл",
    r"собери.*из",
    r"сгенерируй.*из",
    r"создай.*объединив",
    r"сделай.*объединив",
    r"из файла.*и.*файла.*создай",
    r"из файла.*и.*файла.*сделай",
    r"по примеру.*создай",
    r"по шаблону.*создай",
    r"как в примере",
    r"как в examples",
    r"используя.*шаблон",
    r"создай.*excel",
    r"создай.*word",
    r"создай.*xlsx",
    r"создай.*docx",
    r"сделай.*отчёт.*из",
    r"сделай.*отчет.*из",
    r"создай.*отчёт.*из",
    r"создай.*отчет.*из",
    r"создай новый",
    r"сделай новый",
    r"новый файл из",
]


def _is_edit_command(text: str) -> bool:
    text_lower = text.lower()
    for trigger in EDIT_TRIGGERS:
        if re.search(trigger, text_lower):
            return True
    return False


def _is_generate_command(text: str) -> bool:
    text_lower = text.lower()
    for trigger in GENERATE_TRIGGERS:
        if re.search(trigger, text_lower):
            return True
    return False


def _extract_files_from_generate_command(text: str, role: str) -> Tuple[
    List[str], Optional[str], Optional[str], Optional[str]]:
    source_files = []
    output_format = None
    output_name = None
    template = None

    text_lower = text.lower()

    if 'excel' in text_lower or 'xlsx' in text_lower or 'таблиц' in text_lower:
        output_format = 'xlsx'
    elif 'word' in text_lower or 'docx' in text_lower or 'документ' in text_lower:
        output_format = 'docx'
    else:
        output_format = 'xlsx'

    template_match = re.search(r'(?:по примеру|по шаблону|как в|используя шаблон)\s+["\']?([^\s"\']+)["\']?', text,
                               re.I)
    if template_match:
        template = template_match.group(1)

    examples_match = re.search(r'(?:из|из папки)\s+examples?\s*[/\\]?\s*([^\s,]+)', text, re.I)
    if examples_match:
        template = examples_match.group(1)

    output_match = re.search(r'(?:назови|сохрани как|имя файла)\s+["\']?([a-zA-Zа-яА-Я0-9_\-]+)["\']?', text, re.I)
    if output_match:
        output_name = output_match.group(1)

    file_patterns = [
        r'из\s+(?:файлов?\s+)?["\']?([^\s"\']+\.(?:xlsx?|docx|pdf|pptx))["\']?',
        r'(?:файл[аы]?\s+)?["\']?([^\s"\']+\.(?:xlsx?|docx|pdf|pptx))["\']?',
        r'и\s+["\']?([^\s"\']+\.(?:xlsx?|docx|pdf|pptx))["\']?',
    ]

    found_files = set()
    for pattern in file_patterns:
        matches = re.findall(pattern, text, re.I)
        for match in matches:
            if match and match != template:
                found_files.add(match)

    role_dir = STORAGE_DIR / role
    if role_dir.exists():
        for filepath in role_dir.iterdir():
            if filepath.is_file():
                name_lower = filepath.stem.lower()
                if name_lower in text_lower or filepath.name.lower() in text_lower:
                    if filepath.suffix.lower() in ['.xlsx', '.xls', '.docx', '.pdf', '.pptx']:
                        found_files.add(filepath.name)

    source_files = list(found_files)

    return source_files, output_format, output_name, template


def _extract_filename_from_text(text: str, role: str) -> Optional[str]:
    text_lower = text.lower()
    role_dir = get_role_files_dir(role)

    if not role_dir.exists():
        return None

    best_match = None
    best_match_len = 0

    for filepath in role_dir.iterdir():
        if filepath.suffix.lower() in ['.xlsx', '.xls']:
            filename = filepath.name
            filename_lower = filename.lower()

            if filename_lower in text_lower:
                if len(filename) > best_match_len:
                    best_match = filename
                    best_match_len = len(filename)

            stem_lower = filepath.stem.lower()
            if stem_lower in text_lower:
                if len(filepath.stem) > best_match_len:
                    best_match = filename
                    best_match_len = len(filepath.stem)

    if best_match:
        logger.info(f"Найден файл по точному совпадению: {best_match}")
        return best_match

    xlsx_match = re.search(r'(\S+\.xlsx?)', text, re.I)
    if xlsx_match:
        potential_name = xlsx_match.group(1)
        found = _find_file_by_pattern(potential_name, role)
        if found:
            logger.info(f"Найден файл по паттерну: {found}")
            return found

    keywords = []
    for word in text.split():
        word_clean = re.sub(r'[^\w]', '', word.lower())
        if word_clean and len(word_clean) >= 3:
            if word_clean not in ['отредактируй', 'редактируй', 'измени', 'удали',
                                  'добавь', 'файл', 'таблицу', 'таблица', 'excel',
                                  'строку', 'строки', 'колонку', 'ячейку', 'работы',
                                  'все', 'выполненые', 'выполненные', 'невыполненные']:
                keywords.append(word_clean)

    if keywords:
        best_file = None
        best_score = 0

        for filepath in role_dir.iterdir():
            if filepath.suffix.lower() in ['.xlsx', '.xls']:
                stem_lower = filepath.stem.lower()
                score = sum(1 for kw in keywords if kw in stem_lower)
                if score > best_score:
                    best_score = score
                    best_file = filepath.name

        if best_file:
            logger.info(f"Найден файл по ключевым словам ({best_score} совпадений): {best_file}")
            return best_file

    return None


def _find_file_by_pattern(pattern: str, role: str) -> Optional[str]:
    if not pattern:
        return None

    role_dir = get_role_files_dir(role)
    if not role_dir.exists():
        return None

    pattern_clean = pattern.lower().replace('.xlsx', '').replace('.xls', '')
    pattern_clean = re.sub(r'[^\w]', '', pattern_clean)

    best_match = None
    best_score = 0

    for filepath in role_dir.iterdir():
        if filepath.suffix.lower() in ['.xlsx', '.xls']:
            stem_clean = re.sub(r'[^\w]', '', filepath.stem.lower())

            if pattern_clean == stem_clean:
                return filepath.name

            if pattern_clean in stem_clean:
                score = len(pattern_clean) / len(stem_clean)
                if score > best_score:
                    best_score = score
                    best_match = filepath.name

    return best_match


def _is_complex_edit_command(text: str) -> bool:
    complex_patterns = [
        r"удали.*все",
        r"удали.*выполнен",
        r"удали.*невыполнен",
        r"удали.*где",
        r"удали.*которые",
        r"измени.*все",
        r"замени.*все",
        r"пересчитай",
        r"обнови.*итог",
    ]

    text_lower = text.lower()
    for pattern in complex_patterns:
        if re.search(pattern, text_lower):
            return True
    return False


def _get_edit_instruction(text: str, filename: str) -> str:
    text_clean = text.lower()
    text_clean = re.sub(r'отредактируй\s*', '', text_clean)
    text_clean = re.sub(r'редактируй\s*', '', text_clean)
    text_clean = re.sub(r'[^\s]+\.xlsx?', '', text_clean, flags=re.I)
    text_clean = text_clean.strip()
    return text_clean


def _format_content_for_llm(content) -> str:
    parts = []

    if content.text:
        text_preview = content.text[:2000]
        if len(content.text) > 2000:
            text_preview += "\n... (текст обрезан)"
        parts.append(f"Текст:\n{text_preview}")

    for i, table in enumerate(content.tables):
        table_str = f"\nТаблица {i + 1}"
        if table.sheet_name:
            table_str += f" (лист: {table.sheet_name})"
        table_str += ":\n"

        if table.headers:
            table_str += "| " + " | ".join(str(h) for h in table.headers) + " |\n"
            table_str += "| " + " | ".join(["---"] * len(table.headers)) + " |\n"

        for row in table.rows[:50]:
            table_str += "| " + " | ".join(str(cell) for cell in row) + " |\n"

        if len(table.rows) > 50:
            table_str += f"... ещё {len(table.rows) - 50} строк\n"

        parts.append(table_str)

    return "\n".join(parts) if parts else "(пустой файл)"


def _create_file_from_llm_json(data: dict, output_format: str, output_name: str, sources: list, template: str) -> dict:
    from datetime import datetime
    from openpyxl import Workbook
    from openpyxl.styles import Font, Border, Side, Alignment
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        sheets = data.get("sheets", [])
        title = data.get("title", "")

        if output_format in ['xlsx', 'excel', 'xls']:
            wb = Workbook()

            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            header_font = Font(bold=True)

            for sheet_idx, sheet_data in enumerate(sheets):
                if sheet_idx == 0:
                    ws = wb.active
                    ws.title = sheet_data.get("name", "Лист1")[:31]
                else:
                    ws = wb.create_sheet(title=sheet_data.get("name", f"Лист{sheet_idx + 1}")[:31])

                headers = sheet_data.get("headers", [])
                rows = sheet_data.get("rows", [])

                current_row = 1

                if headers:
                    for col_idx, header in enumerate(headers, 1):
                        cell = ws.cell(row=current_row, column=col_idx, value=header)
                        cell.font = header_font
                        cell.border = thin_border
                        cell.alignment = Alignment(horizontal='center')
                    current_row += 1

                for row_data in rows:
                    for col_idx, value in enumerate(row_data, 1):
                        cell = ws.cell(row=current_row, column=col_idx, value=value)
                        cell.border = thin_border
                    current_row += 1

                for col in ws.columns:
                    max_length = 0
                    column = col[0].column_letter
                    for cell in col:
                        try:
                            if cell.value:
                                max_length = max(max_length, len(str(cell.value)))
                        except:
                            pass
                    ws.column_dimensions[column].width = min(max_length + 2, 50)

            filename = f"{output_name}_{timestamp}.xlsx"
            filepath = STORAGE_DIR.parent / "downloads" / filename
            filepath.parent.mkdir(parents=True, exist_ok=True)

            wb.save(filepath)
            wb.close()

            download_url = f"{os.getenv('SERVER_URL', 'http://localhost:8000')}/download/{filename}"

            return {
                "success": True,
                "filename": filename,
                "download_url": download_url,
                "sheets_count": len(sheets)
            }

        elif output_format in ['docx', 'word', 'doc']:
            doc = Document()

            if title:
                heading = doc.add_heading(title, level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for sheet_data in sheets:
                sheet_name = sheet_data.get("name", "")
                if sheet_name:
                    doc.add_heading(sheet_name, level=1)

                headers = sheet_data.get("headers", [])
                rows = sheet_data.get("rows", [])

                if headers or rows:
                    num_cols = len(headers) if headers else len(rows[0]) if rows else 0
                    num_rows = (1 if headers else 0) + len(rows)

                    if num_cols > 0 and num_rows > 0:
                        table = doc.add_table(rows=num_rows, cols=num_cols)
                        table.style = 'Table Grid'

                        if headers:
                            for idx, header in enumerate(headers):
                                cell = table.rows[0].cells[idx]
                                cell.text = str(header)
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.bold = True

                        start_row = 1 if headers else 0
                        for row_idx, row_data in enumerate(rows):
                            for col_idx, value in enumerate(row_data):
                                if col_idx < num_cols:
                                    table.rows[start_row + row_idx].cells[col_idx].text = str(value)

                        doc.add_paragraph()

            filename = f"{output_name}_{timestamp}.docx"
            filepath = STORAGE_DIR.parent / "downloads" / filename
            filepath.parent.mkdir(parents=True, exist_ok=True)

            doc.save(filepath)

            download_url = f"{os.getenv('SERVER_URL', 'http://localhost:8000')}/download/{filename}"

            return {
                "success": True,
                "filename": filename,
                "download_url": download_url,
                "sheets_count": len(sheets)
            }
        else:
            return {"success": False, "error": f"Неподдерживаемый формат: {output_format}"}

    except Exception as e:
        logger.error(f"Ошибка создания файла из JSON: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


async def route_message(messages: list, role: str):
    last_user_msg = messages[-1]["content"]
    state = memory.get_state(role) or {}

    logger.info(f"Router: обрабатываем '{last_user_msg[:50]}...'")

    if state.get("awaiting_file_choice"):
        if state.get("awaiting_excel_choice"):
            from tools.excel_tool import select_excel_file
            chosen_text = select_excel_file(role, last_user_msg)
            state["awaiting_file_choice"] = False
            state["awaiting_excel_choice"] = False
            memory.set_state(role, state)
            return chosen_text, messages
        else:
            chosen_text = select_file(role, last_user_msg)
            state["awaiting_file_choice"] = False
            memory.set_state(role, state)
            return chosen_text, messages

    if state.get("awaiting_file_for_edit"):
        operations = state.get("pending_operations", [])
        filename = _find_file_by_pattern(last_user_msg, role)

        if filename:
            result = edit_excel(filename, operations, role=role)
            state["awaiting_file_for_edit"] = False
            state["pending_operations"] = None
            memory.set_state(role, state)

            if result.get("success"):
                return f"Готово! Скачать: {result['download_url']}", messages
            else:
                return f"Ошибка: {result.get('error')}", messages
        else:
            return "Файл не найден. Укажите точное имя файла.", messages

    gen_match = re.search(
        r'```json\s*(\{[\s\S]*?"sheets"[\s\S]*?\})\s*```',
        last_user_msg,
        re.I
    )
    if gen_match and state.get("pending_file_generation"):
        try:
            gen_data = json.loads(gen_match.group(1))
            pending = state.get("pending_file_generation", {})

            result = _create_file_from_llm_json(
                gen_data,
                pending.get("output_format", "xlsx"),
                pending.get("output_name", "generated"),
                pending.get("sources", []),
                pending.get("template")
            )

            state["pending_file_generation"] = None
            memory.set_state(role, state)

            if result.get("success"):
                response = f"✅ Файл создан: {result['filename']}\n\n"
                response += f"📁 Источники: {', '.join(pending.get('sources', []))}\n"
                response += f"📋 По шаблону: {pending.get('template')}\n"
                response += f"📊 Таблиц: {result.get('sheets_count', 0)}\n\n"
                response += f"🔗 Скачать: {result['download_url']}"
                return response, messages
            else:
                return f"❌ Ошибка создания файла: {result.get('error')}", messages

        except json.JSONDecodeError as e:
            logger.error(f"Ошибка парсинга JSON генерации: {e}")
            state["pending_file_generation"] = None
            memory.set_state(role, state)

    if _is_generate_command(last_user_msg):
        logger.info("Router: обнаружена команда генерации файла")

        if re.search(r'покажи\s+примеры|список\s+примеров|что\s+есть\s+в\s+examples|шаблоны|список шаблонов',
                     last_user_msg, re.I):
            examples = get_example_files()
            if examples:
                examples_list = "\n".join([f"- {e['name']} ({e['type']})" for e in examples])
                return f"📁 Доступные шаблоны в папке examples:\n{examples_list}", messages
            else:
                return "Папка examples пуста. Добавьте шаблоны в storage/examples/", messages

        source_files, output_format, output_name, template = _extract_files_from_generate_command(last_user_msg, role)

        if not source_files:
            role_dir = STORAGE_DIR / role
            available_files = []
            if role_dir.exists():
                available_files = [f.name for f in role_dir.iterdir()
                                   if f.suffix.lower() in ['.xlsx', '.xls', '.docx', '.pdf', '.pptx']]

            if available_files:
                files_list = "\n".join([f"- {f}" for f in available_files[:15]])
                return f"Укажите файлы для объединения.\n\n📁 Доступные файлы:\n{files_list}\n\n💡 Пример: 'Создай Excel из file1.xlsx и file2.docx'", messages
            else:
                return "Файлы не найдены. Загрузите файлы для объединения.", messages

        if not output_name:
            output_name = "combined"

        if template:
            from tools.file_reader_tool import extract_content, read_multiple_files

            template_path = find_file(template, role)
            if not template_path:
                return f"❌ Шаблон не найден: {template}", messages

            template_content = extract_content(template_path)
            source_contents = read_multiple_files(source_files, role)

            if not source_contents:
                return "❌ Не удалось прочитать исходные файлы", messages

            template_preview = _format_content_for_llm(template_content)
            sources_preview = "\n\n---\n\n".join([
                f"ФАЙЛ: {c.filename}\n{_format_content_for_llm(c)}"
                for c in source_contents
            ])

            context = f"""ЗАДАЧА: Создать файл по структуре шаблона, используя данные из исходных файлов.

ШАБЛОН ({template}):
{template_preview}

ИСХОДНЫЕ ДАННЫЕ:
{sources_preview}

---

Проанализируй структуру шаблона и данные из исходных файлов.
Создай JSON который соответствует структуре шаблона, но с данными из исходных файлов.

Формат ответа - ТОЛЬКО JSON:
```json
{{
  "output_format": "{output_format}",
  "output_name": "{output_name}",
  "title": "Заголовок документа",
  "sheets": [
    {{
      "name": "Название листа",
      "headers": ["Колонка1", "Колонка2", ...],
      "rows": [
        ["значение1", "значение2", ...],
        ...
      ]
    }}
  ]
}}
```

Важно:
- Структура должна ТОЧНО соответствовать шаблону
- Данные берутся из исходных файлов
- Объедини похожие данные из разных файлов
- Если в шаблоне несколько разделов (например Материалы и Работы) - сохрани эту структуру
"""
            messages.append({"role": "user", "content": context})

            state["pending_file_generation"] = {
                "output_format": output_format,
                "output_name": output_name,
                "sources": [c.filename for c in source_contents],
                "template": template
            }
            memory.set_state(role, state)

            logger.info(f"Router: генерация по шаблону, отправляем в LLM")
            return None, messages

        title_match = re.search(r'(?:с названием|заголовок|title)\s+["\']?([^"\']+)["\']?', last_user_msg, re.I)
        title = title_match.group(1) if title_match else None

        include_images = 'без картинок' not in last_user_msg.lower() and 'без изображений' not in last_user_msg.lower()

        result = generate_file(
            source_files=source_files,
            output_format=output_format,
            output_name=output_name,
            title=title,
            template_name=None,
            include_images=include_images,
            role=role
        )

        if result.get("success"):
            response = f"✅ Файл создан: {result['filename']}\n\n"
            response += f"📁 Источники: {', '.join(result.get('sources', []))}\n"
            response += f"📊 Таблиц: {result.get('tables_count', 0)}, "
            response += f"Изображений: {result.get('images_count', 0)}\n\n"
            response += f"🔗 Скачать: {result['download_url']}"
            return response, messages
        else:
            return f"❌ Ошибка: {result.get('error')}", messages

    if _is_edit_command(last_user_msg):
        logger.info("Router: обнаружена команда редактирования")
        filename = _extract_filename_from_text(last_user_msg, role)
        logger.info(f"Router: извлечённый файл = {filename}")

        if not filename:
            results = smart_search(last_user_msg, role, limit=5)
            excel_files = [r for r in results if r.get("is_table")]

            if len(excel_files) == 1:
                filename = excel_files[0]["filename"]
            elif len(excel_files) > 1:
                files_list = "\n".join([f"- {f['filename']}" for f in excel_files])
                return f"Найдено несколько файлов:\n{files_list}\n\nУкажите какой файл редактировать.", messages

        if not filename:
            role_dir = get_role_files_dir(role)
            all_excel = []
            if role_dir.exists():
                all_excel = [f.name for f in role_dir.iterdir()
                             if f.suffix.lower() in ['.xlsx', '.xls']]
            if all_excel:
                files_list = "\n".join([f"- {f}" for f in all_excel[:10]])
                return f"Не удалось определить файл. Доступные файлы:\n{files_list}", messages
            else:
                return "Excel файлы не найдены.", messages

        if _is_complex_edit_command(last_user_msg):
            instruction = _get_edit_instruction(last_user_msg, filename)
            file_content = read_excel_for_edit(filename, role=role)

            context = f"""Файл: {filename}

ВАЖНО: Колонка ROW содержит РЕАЛЬНЫЕ номера строк Excel. Используй именно эти номера в операциях!

Содержимое:
{file_content}

---
Инструкция пользователя: {instruction}

Проанализируй таблицу и сгенерируй JSON с операциями редактирования.
Формат ответа:
```json
{{
  "filename": "{filename}",
  "operations": [
    {{"action": "delete_row", "row": N}},
    ...
  ]
}}
```

Доступные операции:
- delete_row: удалить строку (row = номер строки)
- edit_cell: изменить ячейку (row, col, value)
- add_row: добавить строку (data = массив значений, after_row = после какой строки)
"""
            messages.append({"role": "user", "content": context})
            logger.info(f"Router: сложная команда, отправляем в LLM. Файл: {filename}")
            return None, messages

        _, operations = parse_excel_command(last_user_msg)

        if operations:
            result = edit_excel(filename, operations, role=role)

            if result.get("success"):
                ops_desc = ", ".join([op["action"] for op in operations])
                return f"Выполнено ({ops_desc})!\n\nСкачать: {result['download_url']}", messages
            else:
                return f"Ошибка: {result.get('error')}", messages
        else:
            file_content = read_excel_for_edit(filename, role=role)
            instruction = _get_edit_instruction(last_user_msg, filename)

            context = f"""Файл: {filename}

ВАЖНО: Колонка ROW содержит РЕАЛЬНЫЕ номера строк Excel. Используй именно эти номера в операциях!

Содержимое:
{file_content}

---
Инструкция: {instruction}

Сгенерируй JSON с операциями:
```json
{{
  "filename": "{filename}",
  "operations": [...]
}}
```
"""
            messages.append({"role": "user", "content": context})
            return None, messages

    if re.search(r"(найди|поиск|найди в файлах|search)\s+\w", last_user_msg, re.I):
        query = re.sub(r"(найди|поиск|найди в файлах|search)\s*", "", last_user_msg, flags=re.I).strip()
        if query:
            result = perform_search(query, role)
            return result, messages

    if re.search(r"(запомни|сохрани факт|добавь в память)", last_user_msg, re.I):
        fact = re.sub(r"(запомни|сохрани факт|добавь в память)\s*", "", last_user_msg, flags=re.I).strip()
        if fact and vector_store.is_connected():
            result = vector_store.add_memory(fact, "general", role)
            return result.get("message", "Ошибка"), messages

    if re.search(r"(сводка|сводку|обзор|summary).*(файл|документ|всех)", last_user_msg, re.I):
        result = await process_multiple_files(last_user_msg, role, top_n=20)
        return result, messages

    if re.search(r"(сравни|сравнение|compare)", last_user_msg, re.I):
        result = await process_multiple_files(last_user_msg, role, top_n=10)
        return result, messages

    edit_match = re.search(
        r'```json\s*(\{[\s\S]*?"operations"[\s\S]*?\})\s*```',
        last_user_msg,
        re.I
    )
    if edit_match:
        try:
            edit_data = json.loads(edit_match.group(1))
            filename = edit_data.get("filename")
            operations = edit_data.get("operations", [])

            if filename and operations:
                result = edit_excel(filename, operations, role=role)
                if result.get("success"):
                    return f"Файл отредактирован!\n\nСкачать: {result['download_url']}", messages
                else:
                    return f"Ошибка: {result.get('error')}", messages
        except json.JSONDecodeError as e:
            logger.error(f"Ошибка парсинга JSON: {e}")

    if any(ext in last_user_msg.lower() for ext in [".xlsx", ".xls"]):
        filename = _extract_filename_from_text(last_user_msg, role)
        if filename:
            content = read_excel(filename, role=role)
            return content, messages

    file_result = try_handle_file_command(last_user_msg, role)
    if file_result:
        return file_result, messages

    return None, messages